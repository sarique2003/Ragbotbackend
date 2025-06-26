import logging
from typing import List, Tuple
from pathlib import Path
import pandas as pd
import asyncio
from dotenv import load_dotenv
from backend.helpers import get_env_value

load_dotenv()
import pdfplumber
import camelot
import re
import docx
from backend.services.rag_service.base_vector_db import VectorDB
from backend.models.vectordb_models import PineconeConfig
from backend.services.rag_service.pinecone import PineconeDB
from langchain.docstore.document import Document

logger = logging.getLogger(__name__)

class VectorDbService:
    def __init__(self, vector_db: VectorDB):
        self.vector_db = vector_db
        self.namespace = "insurance_namespace"

    # ───────────────────────── helpers ──────────────────────────
    @staticmethod
    def _row_to_text(row: pd.Series) -> str:
        return "\n".join(
            f"{col} is {row[col]}" for col in row.index if str(row[col]).strip()
        )

    @staticmethod
    def _split_paragraphs(text: str) -> List[str]:
        paras = [
            p.strip()
            for p in re.split(r"\n\s*\n", text)
            if p and len(p.strip()) > 25  # heuristically keep “real” paragraphs
        ]
        return paras

    # ---------------------- PDF extraction ----------------------
    @staticmethod
    def _tables_from_pdf(path: Path) -> List[pd.DataFrame]:
        try:
            tables = camelot.read_pdf(str(path), pages="all", strip_text="\n")
            return [t.df for t in tables] if tables else []
        except Exception as e:
            logger.debug(f"Camelot failed on {path.name}: {e}")

        dfs: List[pd.DataFrame] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                tbl = page.extract_table()
                if tbl and len(tbl) > 1:
                    dfs.append(pd.DataFrame(tbl[1:], columns=tbl[0]))
        return dfs

    @staticmethod
    def _paragraphs_from_pdf(path: Path) -> List[str]:
        paras: List[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                paras.extend(VectorDbService._split_paragraphs(text))
        return paras

    @staticmethod
    def _tables_from_docx(path: Path) -> List[pd.DataFrame]:
        doc = docx.Document(str(path))
        dfs: List[pd.DataFrame] = []
        for t in doc.tables:
            rows = [[c.text.strip() for c in r.cells] for r in t.rows]
            if len(rows) > 1:
                dfs.append(pd.DataFrame(rows[1:], columns=rows[0]))
        return dfs

    @staticmethod
    def _paragraphs_from_docx(path: Path) -> List[str]:
        doc = docx.Document(str(path))
        return [p.text.strip() for p in doc.paragraphs if len(p.text.strip()) > 25]


    async def ingest_to_vector_db(self, folder_path: str):
        folder = Path(folder_path).expanduser().resolve()
        if not folder.is_dir():
            raise FileNotFoundError(f"{folder} is not a directory")

        texts: List[str] = []
        for file in folder.iterdir():
            if not file.is_file():
                continue

            try:
                if file.suffix.lower() == ".pdf":
                    # --- body text ---
                    paras = self._paragraphs_from_pdf(file)
                    texts.extend(paras)
                    # --- tables ---
                    for df in self._tables_from_pdf(file):
                        texts.extend(self._row_to_text(row) for _, row in df.iterrows())
                    logger.info(f"{file.name}: added {len(paras)} paragraphs + tables")

                elif file.suffix.lower() == ".docx":
                    paras = self._paragraphs_from_docx(file)
                    texts.extend(paras)
                    for df in self._tables_from_docx(file):
                        texts.extend(self._row_to_text(row) for _, row in df.iterrows())
                    logger.info(f"{file.name}: added {len(paras)} paragraphs + tables")

            except Exception as exc:
                logger.warning(f"{file.name} skipped: {exc}")

        if not texts:
            logger.warning("Nothing to ingest – no paragraphs or tables detected.")
            return

        ids = await self.vector_db.ingest_data(texts, namespace=self.namespace)
        logger.info(f"Ingested {len(ids)} docs (paras + rows) into namespace '{self.namespace}'")

    async def retrieve(self, query: str) -> List[Tuple[Document, float]]:
        return await self.vector_db.similarity_search_with_score(query=query, namespace=self.namespace)



async def main():
    # Assuming you have a DataFrame ready (replace with actual DataFrame)
    data = {
        "title": ["Hair Oil", "Hair Serum"],
        "description": ["Nourishing oil for hair growth", "Smoothening serum for frizzy hair"],
        "price": [25.00, 30.00]
    }
    dataframe = pd.DataFrame(data)
    #namespace = "hair_originals_index_06_01_2025"


    config = PineconeConfig(
        api_key=get_env_value('PINECONE_API_KEY'),
        base_url=get_env_value('PINECONE_URL')
    )
    pineconedb = PineconeDB(config=config)
    vectorservice = VectorDbService(vector_db= pineconedb)
    index = "sarique-test"
    api_key = "pcsk_5WohRM_EncoaGWkcYDnHno1FymNXXiKwcKe9c9JRBHUDk2M7gAMRHcxuMYaX2Vuzx2mnML"
    api_environment = "gcp-starter"
    await vectorservice.ingest_to_vector_db('/Users/mohdsarique/Documents/Chabot/backend/data')

if __name__ == '__main__':
    asyncio.run(main())